## ---------------------------------------------------------------------------##
# Author: Beatrix Haddock
# Date: 2024-08-29
# Purpose:
# Pull and export to xlsx:
#   - 'CAP Version #'
#   - 'Last CAP revision Date'
#   - CAP Version # (DD-MM-YYYY) last distributed
#   - CAP distributed Description
## ---------------------------------------------------------------------------##
import pandas as pd
import numpy as np
import datetime
import datefinder
import docx
import io
import re

from office365.runtime.auth.authentication_context import AuthenticationContext
from office365.sharepoint.client_context import ClientContext
from office365.sharepoint.files.file import File

import yaml
import sys
import os
## ---------------------------------------------------------------------------##
yaml_path = "/home/bhaddock/repos/sdmc_cap_tracker/config.yaml"
with open(yaml_path, 'r') as file:
    config = yaml.safe_load(file)

username_shrpt = 'bhaddock@fredhutch.org'
password_shrpt = config['password']

def get_sharepoint_data():
    # pull CAPs --------------------------------------------------------------##
    old_stdout = sys.stdout # backup current stdout
    sys.stdout = open(os.devnull, "w")
    CAP_DOCS, cap_links = compile_CAP_docs_and_fnames()
    sys.stdout = old_stdout # reset old stdout

    # useful for keying into CAP_DOCS
    cap_links['name'] = cap_links.network + cap_links.protocol

    # parse out CAP version and date of last update ---------------------------##
    cap_tracking = cap_links.drop(columns=['url','folder'])
    cap_tracking['header'] = cap_tracking.name.apply(lambda x: pull_header(x, CAP_DOCS))
    cap_tracking = parse_version_and_date(cap_tracking)
    cap_tracking = cap_tracking.sort_values(by=['network','protocol'])

    # report on issues -------------------------------------------------------##
    issues = cap_tracking.loc[(
        (cap_tracking.cap_version.str.contains(" OR |No version number found")) |
        (cap_tracking.last_cap_revision_date.str.contains("Error|NA"))
    ), ['network','protocol','cap_version','last_cap_revision_date']]

    print("\nChecking for CAP updates; unable to resolve the following:")
    print(issues)

    # deprecated -- they didnt actually want this ----------------------------##
    # pull protocol version in CAP
    # last_tables = {}
    # for k in CAP_DOCS.keys():
    #     try:
    #         last_tables[k] = get_nth_table_as_list(CAP_DOCS[k], i=-1)
    #     except:
    #         last_tables[k] = "ERROR"
    #
    # protocol_version_map = {p: get_protocol_version_from_last_table(last_tables[p]) for p in last_tables.keys()}
    # cap_tracking['most_recent_protocol_version_in_CAP_modification_log'] = (cap_links.network + cap_links.protocol).map(protocol_version_map)

    # pull last distributed CAP info -----------------------------------------##
    first_tables = {}
    for k in CAP_DOCS.keys():
        try:
            first_tables[k] = get_nth_table_as_list(CAP_DOCS[k], i=0)
        except:
            first_tables[k] = "Error pulling first table"

    cap_tracking["cap_version_last_distributed"] = cap_tracking.name.apply(
        lambda x: find_last_distributed_CAP_info(x, first_tables)[0]
    )
    cap_tracking["cap_last_distributed_description"] = cap_tracking.name.apply(
        lambda x: find_last_distributed_CAP_info(x, first_tables)[1]
    )

    # return table -----------------------------------------------------------##
    return cap_tracking


    # # check diffs ------------------------------------------------------------##
    # last_week = (datetime.datetime.now() - datetime.timedelta(days=7)).date().isoformat()
    # try:
    #     last = pd.read_excel(savedir + f"CAP_versions_and_dates_{last_week}.xlsx")
    # except:
    #     last = [i for i in os.listdir(savedir) if "CAP_versions_and_dates" in i]
    #     last = np.sort(last)[-1]
    #     print(f"Using {last} because couldn't find one from exactly {last_week}")
    #     last = pd.read_excel(savedir + last)
    #
    # # convert types for merge
    # last.protocol = last.protocol.astype(str)
    # cap_tracking.protocol = cap_tracking.protocol.astype(str)
    #
    # # merge
    # diff = cap_tracking[['network','protocol','presumed_date','presumed_version']].merge(
    #     last,
    #     on=['network','protocol'],
    #     how='outer',
    #     suffixes=('_new', '_old'),
    #     )
    #
    # # subset to rows with changes, columns of interest
    # usecols = [
    #     'network',
    #     'protocol',
    #     'presumed_date_old',
    #     'presumed_date_new',
    #     'presumed_version_old',
    #     'presumed_version_new'
    #     ]
    # diff = diff.loc[
    #         (diff.presumed_date_old!=diff.presumed_date_new) | (diff.presumed_version_old!=diff.presumed_version_new),
    #          usecols]
    #
    # # if changes, report
    # if len(diff) > 0:
    #     # diff.to_excel(savedir + f"updates_{today}.xlsx", index=False)
    #     diff.to_excel(usedir + f"updates_{today}.xlsx", index=False)
    #     print(f"\nChanges in the following rows:")
    #     print(diff)


## HELPERS - pull CAPS from sharepoint ---------------------------------------##
def compile_CAP_docs_and_fnames():
    cap_links = pd.read_csv("/home/bhaddock/repos/sdmc_cap_tracker/cap_sharepoint_links.txt", sep="\t")
    cap_links['filename'] = 'na'

    CAP_DOCS = {}
    for i, row, in cap_links.iterrows():
        doc, fname = get_doc_and_filename(row.url, row.folder, row.protocol)
        CAP_DOCS[row.network + str(row.protocol)] = doc
        cap_links.loc[cap_links.protocol==row.protocol,'filename'] = fname

    return CAP_DOCS, cap_links

def get_doc_and_filename(url, folder, protocol):
    ###Authentication###For authenticating into your sharepoint site###
    ctx_auth = AuthenticationContext(url)
    if ctx_auth.acquire_token_for_user(username_shrpt, password_shrpt):
      ctx = ClientContext(url, ctx_auth)
      web = ctx.web
      ctx.load(web)
      ctx.execute_query()
      print('Authenticated into sharepoint as: ',web.properties['Title'])

    else:
      print(ctx_auth.get_last_error())

    def print_folder_contents(ctx, folder_url):
        try:
            folder = ctx.web.get_folder_by_server_relative_url(folder_url)
            fold_names = []
            sub_folders = folder.files #Replace files with folders for getting list of folders
            ctx.load(sub_folders)
            ctx.execute_query()

            for s_folder in sub_folders:

                fold_names.append(s_folder.properties["Name"])

            return fold_names

        except Exception as e:
            print('Problem printing out library contents: ', e)

    filenames = print_folder_contents(ctx, folder)
    protocol_number = re.sub('[^0-9]','', protocol)
    fnames = [i for i in filenames if 'cap' in i.lower() and 'doc' in i.lower() and protocol_number in i.lower()]
    if len(fnames) == 0:
        print(f"{protocol }CAP not found. Filenames in CAP folder: {filenames}'")
        return "No files found", "NA"
    elif len(fnames) > 1:
        print(f"multiple hits for matching filenames: {fnames}. Taking {np.sort(fnames)[-1]}")
    fname = np.sort(fnames)[-1]
    file_url_shrpt = folder + "/" + fname

    #Load the sharepoint file content to "response" variable
    response = File.open_binary(ctx, file_url_shrpt)

    #save data to BytesIO stream
    bytes_file_obj = io.BytesIO()
    bytes_file_obj.write(response.content)
    bytes_file_obj.seek(0) #set file object to start
    doc = docx.Document(bytes_file_obj)

    return doc, fname

## HELPERS - pulling CAP version and last updated date -----------------------##
def pull_header(network_protocol, CAP_DOCS):
    doc = CAP_DOCS[network_protocol]
    if isinstance(doc, str):
        return doc
    header = doc.sections[0].header
    header_text = '\t'.join([i.text for i in header.paragraphs])
    return header_text

def parse_version_and_date(cap_tracking):
    cap_tracking.header = cap_tracking.header.str.replace("\t\t\t","\t")
    cap_tracking.header = cap_tracking.header.str.replace("\t\t","\t")

    cap_tracking['version_and_date_from_header'] = cap_tracking.header.str.split("\t", expand=True)[1]

    cap_tracking['version_from_filename'] = cap_tracking.filename.apply(find_version_number_from_text)
    cap_tracking['version_from_header'] = cap_tracking.version_and_date_from_header.apply(find_version_number_from_header)

    cap_tracking['date_from_filename'] = cap_tracking.filename.apply(get_date_from_filename)
    cap_tracking['date_from_header'] = cap_tracking.version_and_date_from_header.apply(get_date_from_text)

    cap_tracking['cap_version'] = cap_tracking.apply(lambda x: get_presumed_version(x.version_from_header, x.version_from_filename), axis=1)
    cap_tracking['last_cap_revision_date'] = cap_tracking.apply(lambda x: get_presumed_date(x.date_from_header, x.date_from_filename), axis=1)

    cols = [
        'network',
        'protocol',
        'cap_folder_sharepoint_path',
        'cap_version',
        'last_cap_revision_date'
    ]
    if 'name' in cap_tracking.columns:
        cols = ['name'] + cols
    return cap_tracking[cols]

def find_version_number_from_text(fname):
    """
    Given an input string (filename)
    returns the first substring that is formatted as {int(s)}.{int(s)}
    """
    try:
        decimal_idx = fname.index(".")
    except:
        return "No version number found"
    before = decimal_idx
    while fname[before - 1].isdigit():
        before -= 1
    # if the first one failed, then the char before the . is not numeric; check if there are later .s
    if before == decimal_idx:
        return find_version_number_from_text(fname[decimal_idx + 1:])

    after = decimal_idx
    while fname[after + 1].isdigit():
        after += 1
    if after == decimal_idx:
        return find_version_number_from_text(fname[decimal_idx + 1:])

    return "V" + fname[before:after + 1]

def find_version_number_from_header(header):
    if header is None:
        return "No version number found"
    v_from_header = find_version_number_from_text(header)
    if v_from_header=="No version number found" and "Draft" in header:
        return "Draft"
    elif "--" in header:
        return "Trouble parsing due to suggested edits"
    else:
        return v_from_header

def extract_date(h, strict=True):
    try:
        dates = list(datefinder.find_dates(h, strict=strict))
        if len(dates) > 0:
            return dates[0].date().isoformat()
        else:
            return "NA"
    except:
        return "Error"

def get_date_from_text(text):
    version = find_version_number_from_text(text)
    if version != "No version number found":
        text = text.replace(version[1:], "")
    return extract_date(text, strict=False)

def get_date_from_filename(text):
    version = find_version_number_from_text(text)
    if version != "No version number found":
        text = text.replace(version[1:], "")
    return extract_date(text, strict=True)

def get_presumed_version(from_header, from_filename):
    if from_header==from_filename:
        return from_header
    if from_filename=="No version number found":
        return from_header
    elif from_filename!="No version number found":
        if from_header=="Draft":
            return from_filename
        elif from_header=="Trouble parsing due to suggested edits":
            return from_filename
        else:
            return f"{from_header} OR {from_filename}?"

def get_presumed_date(from_header, from_filename):
    if from_header==from_filename:
        return from_header
    if from_header=="NA" and from_filename!="NA":
        return from_filename
    if from_header!="NA" and from_filename=="NA":
        return from_header
    if from_header!="NA" and from_filename!="NA":
        return f"{from_header} OR {from_filename}?"

## HELPERS - pulling protocol version from CAP -------------------------------##
def get_nth_table_as_list(doc, i=0):
    table = doc.tables[i]
    l = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text:
                l[i][j] = cell.text
    return l

def get_protocol_version_from_last_table(table):
    i = -1
    j = len(table)
    found = 0
    while found == 0:
        if -1*i > j:
           return "Couldn't find version"
        row = table[i]
        seek = [i for i in row if 'protocol version' in i.lower()]
        if len(seek) > 0:
            version = np.unique(seek)[0]
            version = version.replace("\n","")
            idx = version.lower().index("protocol version")
            version = version[idx + len("protocol version"):]
            if len(version)>0:
                if version[0] == ":":
                    version = version[1:]
            else:
                return "Document being edited"
            version = version.strip()
            if len(version) > 0:
                if version[0].isdigit():
                    version = "V" + version
                return version
            else:
                return "Document being edited"
        else:
            i -= 1
## HELPERS - pulling last distributed CAP version/date/description -----------##
def find_last_distributed_CAP_info(network_protocol, first_tables):
    table = first_tables[network_protocol]
    if isinstance(table, str):
        return table, "NA"
    table = pd.DataFrame(table[1:], columns=table[0])

    cap_version_col = [i for i in table.columns if "cap version" in i.lower()]
    if len(cap_version_col) > 0:
        cap_version_col = cap_version_col[0]
    else:
        return "First table doesn't contain CAP versioning info", "NA"

    versions = table[cap_version_col]
    if isinstance(versions, str):
        return versions, "NA"
    else:
        versions = [i for i in versions if str(i).upper() not in ['NA', 'N/A', 'NAN', '']]
        if len(versions) > 0:
            version = versions[-1]
            description = table.loc[table[cap_version_col]==version].iloc[:,2].values[0]
            return versions[-1].strip(), description
        else:
            return "No version found", "NA"


if __name__=="__main__":
    cap_tracking = get_sharepoint_data()

    savedir = '/networks/vtn/lab/SDMC_labscience/operations/projects/CAP_projectfiles/project_management_ideas/'
    today = datetime.date.today().isoformat()
    cap_tracking.to_excel(savedir + f"EXAMPLE_of_automated_CAP_columns_{today}.xlsx")
